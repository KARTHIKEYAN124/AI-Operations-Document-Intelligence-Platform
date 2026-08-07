from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO

from docx import Document as DocxDocument
from fastapi import UploadFile
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
    "application/rtf",
    "text/rtf",
    "image/png",
    "image/jpeg",
    "image/tiff",
    "image/webp",
    "image/bmp",
}
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024


@dataclass(frozen=True)
class UploadValidation:
    accepted: bool
    reason: str | None
    content_hash: str | None = None
    size_bytes: int = 0
    content: bytes = b""


async def validate_upload(file: UploadFile) -> UploadValidation:
    content = await file.read()
    await file.seek(0)

    if not is_allowed(file.filename or "", file.content_type or ""):
        return UploadValidation(False, "Unsupported file type")
    if len(content) > MAX_FILE_SIZE_BYTES:
        return UploadValidation(False, "File exceeds maximum size")
    if len(content) == 0:
        return UploadValidation(False, "Document is empty")

    try:
        extract_text(file.filename or "upload", content, file.content_type or "")
    except ValueError as exc:
        return UploadValidation(False, str(exc), size_bytes=len(content), content=content)

    return UploadValidation(True, None, sha256(content).hexdigest(), len(content), content)


def is_allowed(filename: str, mime_type: str) -> bool:
    return mime_type in ALLOWED_MIME_TYPES or filename.lower().endswith((".pdf", ".doc", ".docx", ".txt", ".md", ".markdown", ".rtf", ".csv", ".json", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp"))


def extract_text(filename: str, content: bytes, mime_type: str = "") -> tuple[str, list[str], int | None]:
    if not content:
        raise ValueError("Document is empty")

    lower_name = filename.lower()
    if mime_type in {"text/plain", "text/markdown", "text/csv", "application/json", "application/rtf", "text/rtf"} or lower_name.endswith((".txt", ".md", ".markdown", ".rtf", ".csv", ".json")):
        return content.decode("utf-8", errors="replace").strip(), [], None

    if mime_type == "application/pdf" or lower_name.endswith(".pdf"):
        return extract_pdf_text(content)

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or lower_name.endswith(".docx"):
        return extract_docx_text(content)

    if mime_type == "application/msword" or lower_name.endswith(".doc"):
        return extract_legacy_doc_text(content)

    if mime_type.startswith("image/") or lower_name.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp")):
        return extract_image_text(content)

    raise ValueError("Unsupported file type")


def extract_pdf_text(content: bytes) -> tuple[str, list[str], int | None]:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise ValueError("PDF appears to be corrupted") from exc

    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {index}]\n{text.strip()}")

    limitations = []
    if not pages:
        limitations.append("No selectable PDF text was found. OCR is required for scanned pages.")
    return "\n\n".join(pages).strip(), limitations, len(reader.pages)


def extract_docx_text(content: bytes) -> tuple[str, list[str], int | None]:
    try:
        document = DocxDocument(BytesIO(content))
    except Exception as exc:
        raise ValueError("DOCX appears to be corrupted") from exc
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    return "\n".join([*paragraphs, *table_cells]).strip(), [], None


def extract_legacy_doc_text(content: bytes) -> tuple[str, list[str], int | None]:
    text = content.decode("latin1", errors="ignore")
    clean = "".join(character if character in "\n\r\t" or 32 <= ord(character) <= 126 else " " for character in text)
    clean = " ".join(clean.split())
    return clean, ["Used best-effort legacy DOC text extraction. Convert to DOCX for higher fidelity."], None


def extract_image_text(content: bytes) -> tuple[str, list[str], int | None]:
    try:
        image = Image.open(BytesIO(content))
        image.verify()
    except UnidentifiedImageError as exc:
        raise ValueError("Image appears to be corrupted") from exc

    try:
        import pytesseract

        image = Image.open(BytesIO(content))
        text = pytesseract.image_to_string(image).strip()
        return text, ["OCR accuracy depends on image quality."] if text else ["OCR found no readable text."], 1
    except Exception:
        return "", ["OCR adapter is not installed on this backend. Browser OCR or a Tesseract-enabled worker is required."], 1
